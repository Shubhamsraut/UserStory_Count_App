import re
import io
import pandas as pd
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
import streamlit as st

# ------------------------------
# Parsing helpers
# ------------------------------

def iter_block_items(parent):
    xml_children = list(parent.element.body.iterchildren())
    table_iter_index = 0
    for idx, child in enumerate(xml_children):
        if isinstance(child, CT_P):
            p_count = sum(1 for c in xml_children[: idx + 1] if isinstance(c, CT_P))
            yield ("p", parent.paragraphs[p_count - 1])
        elif isinstance(child, CT_Tbl):
            if table_iter_index < len(parent.tables):
                yield ("t", parent.tables[table_iter_index])
            table_iter_index += 1

EPIC_RE = re.compile(r"^\s*Epic\s+(\d+)\s*[:\-\u2013\u2014]\s*(.+)\s*$", re.IGNORECASE)
STORY_RE = re.compile(r"^\s*(?:User\s+)?Story\s+(\d+(?:\.\d+)*)\s*[:\-\u2013\u2014]\s*(.+)\s*$", re.IGNORECASE)

# Header canonicalization
HEADER_ALIASES = {
    "scenario": "Scenario",
    "given": "Given",
    "precondition": "Given",
    "when": "When",
    "action": "When",
    "then": "Then",
    "expected": "Expected",
    "expected result": "Expected",
    "result": "Expected",
    "acceptance criteria": "Acceptance Criteria",
    "criteria": "Acceptance Criteria",
    "ac": "Acceptance Criteria",
    "#": "AC #",
    "no": "AC #",
    "id": "AC #",
    "sr no": "AC #",
    "s no": "AC #",
    "sno": "AC #",
    "srno": "AC #",
    "ac #": "AC #",
    "ac no": "AC #",
    "ac number": "AC #",
}

def _canon_header(text: str) -> str:
    """
    Normalize header text by lowercasing, removing punctuation (except '#'),
    and collapsing whitespace before alias lookup.
    Handles variants like 'Sr. No', 'S.No', 'Sr.No.' etc.
    """
    raw = (text or "").strip()
    t = raw.lower()
    t = re.sub(r"[^\w#]+", " ", t)   # keep word chars and '#', turn punctuation into spaces
    t = re.sub(r"\s+", " ", t).strip()
    return HEADER_ALIASES.get(t, raw if raw else "")

def looks_like_ac_table(table):
    """Heuristic: treat a table as AC if first/second row mentions AC keywords."""
    keywords = {"acceptance", "criteria", "scenario", "given", "when", "then", "expected", "result"}
    for hdr_idx in (0, 1):
        if hdr_idx >= len(table.rows):
            break
        cells_text = " | ".join(cell.text.strip().lower() for cell in table.rows[hdr_idx].cells)
        if any(k in cells_text for k in keywords):
            return True, hdr_idx
    return False, None

def _row_is_empty(row) -> bool:
    return all(not (cell.text or "").strip() for cell in row.cells)

def _first_data_row_index(table, header_row_index):
    total_rows = len(table.rows)
    data_start = header_row_index + 1
    # Skip one possible blank row after header
    if data_start < total_rows and _row_is_empty(table.rows[data_start]):
        data_start += 1
    return min(data_start, total_rows)

def count_ac_rows(table, header_row_index):
    total_rows = len(table.rows)
    data_start = _first_data_row_index(table, header_row_index)
    count = 0
    for r in range(data_start, total_rows):
        if not _row_is_empty(table.rows[r]):
            count += 1
    return max(count, 0)

def parse_ac_table_rows_minimal(table, header_row_index):
    """
    Parse AC table rows into minimal dicts with just 'AC #' and 'Scenario'.
    Robustly identifies the AC number column even if header isn't a perfect match
    (e.g., 'Sr. No', 'S.No', 'Sr.No.'). If only a single 'Acceptance Criteria'
    column exists, maps that to 'Scenario'.
    """
    header_cells = table.rows[header_row_index].cells
    headers = [_canon_header(c.text) for c in header_cells]

    # Figure out best-effort indices
    def _find_idx(names):
        for i, h in enumerate(headers):
            if h in names:
                return i
        return None

    idx_acnum = _find_idx({"AC #"})
    idx_scenario = _find_idx({"Scenario"})
    idx_free_ac = _find_idx({"Acceptance Criteria"})  # fallback to free-text AC column

    data_start = _first_data_row_index(table, header_row_index)

    # Fallback: infer AC # column by value pattern like 2.1.1, 1.2, 3, etc.
    if idx_acnum is None and len(table.rows) > data_start:
        numberlike = re.compile(r"^\s*\d+(?:\.\d+)*\s*$")
        max_sample = min(len(table.rows), data_start + 6)
        best_col, best_hits = None, -1
        num_cols = len(header_cells)
        for col in range(num_cols):
            hits = 0
            for r in range(data_start, max_sample):
                cells = table.rows[r].cells
                if col >= len(cells):
                    continue
                val = (cells[col].text or "").strip()
                if val and numberlike.match(val):
                    hits += 1
            if hits > best_hits:
                best_hits, best_col = hits, col
        # if at least 2 matches, assume this is the AC # column
        if best_hits >= 2:
            idx_acnum = best_col

    out = []
    for r_idx in range(data_start, len(table.rows)):
        row = table.rows[r_idx]
        if _row_is_empty(row):
            continue
        cells = [cell.text.strip() for cell in row.cells]

        ac_no = ""
        if idx_acnum is not None and idx_acnum < len(cells):
            ac_no = cells[idx_acnum].strip()

        scenario = ""
        if idx_scenario is not None and idx_scenario < len(cells):
            scenario = cells[idx_scenario].strip()
        elif idx_free_ac is not None and idx_free_ac < len(cells):
            scenario = cells[idx_free_ac].strip()

        if ac_no or scenario:
            out.append({"AC #": ac_no, "Scenario": scenario})
    return out

def extract_user_stories_and_acs(docx_file):
    """
    Returns:
      stories_df: Module, Epic, Story ID, Story Title, Acceptance Criteria Count
      ac_df (minimal): Module, Epic, Story ID, Story Title, AC #, Scenario
    """
    doc = Document(docx_file)
    paragraphs = doc.paragraphs
    if len(paragraphs) == 0:
        stories_cols = ["Module", "Epic", "Story ID", "Story Title", "Acceptance Criteria Count"]
        ac_cols = ["Module","Epic","Story ID","Story Title","AC #","Scenario"]
        return pd.DataFrame(columns=stories_cols), pd.DataFrame(columns=ac_cols)

    paragraphs_nonempty = [p.text.strip() for p in paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs_nonempty)

    module_match = re.search(r"Module\s*[:\-\u2013\u2014]\s*(.+)", full_text, flags=re.IGNORECASE)
    module = module_match.group(1).strip() if module_match else "Unknown"

    stories = []
    ac_rows = []
    current_epic = None
    current_story = None

    for kind, obj in iter_block_items(doc):
        if kind == "p":
            line = (obj.text or "").strip()
            if not line:
                continue

            em = EPIC_RE.match(line)
            if em:
                epic_num, epic_title = em.groups()
                current_epic = f"{epic_num}: {epic_title.strip()}"
                continue

            sm = STORY_RE.match(line)
            if sm:
                story_id, story_title = sm.groups()
                current_story = {
                    "Module": module,
                    "Epic": current_epic or "Unknown",
                    "Story ID": story_id.strip(),
                    "Story Title": story_title.strip(),
                    "Acceptance Criteria Count": 0
                }
                stories.append(current_story)
                continue

        elif kind == "t":
            if current_story is None:
                continue
            is_ac, hdr_idx = looks_like_ac_table(obj)
            if not is_ac:
                continue

            # Count & parse
            row_count = count_ac_rows(obj, hdr_idx)
            current_story["Acceptance Criteria Count"] += row_count

            parsed = parse_ac_table_rows_minimal(obj, hdr_idx)
            for entry in parsed:
                ac_rows.append({
                    "Module": current_story["Module"],
                    "Epic": current_story["Epic"],
                    "Story ID": current_story["Story ID"],
                    "Story Title": current_story["Story Title"],
                    **entry
                })

    stories_cols = ["Module", "Epic", "Story ID", "Story Title", "Acceptance Criteria Count"]
    ac_cols = ["Module","Epic","Story ID","Story Title","AC #","Scenario"]

    stories_df = pd.DataFrame(stories, columns=stories_cols) if stories else pd.DataFrame(columns=stories_cols)
    ac_df = pd.DataFrame(ac_rows, columns=ac_cols) if ac_rows else pd.DataFrame(columns=ac_cols)

    return stories_df, ac_df

# ------------------------------
# Streamlit UI
# ------------------------------
st.set_page_config(page_title="User Story Extractor", page_icon="📄", layout="wide")

st.markdown(
    """
    <style>
    .block-container {
        border: 2px solid #e0e0e0;
        padding: 2rem;
        border-radius: 10px;
        background-color: #f9fafb;
    }
    .section-title {
        font-size: 1.2rem;
        font-weight: 600;
        padding: 0.5rem;
        margin-top: 2rem;
        border: 1px solid #ddd;
        border-radius: 6px;
        background: #f0f0f5;
    }
    .metric-card {
        text-align: center;
        padding: 0.8rem 1rem;
        border-radius: 10px;
        border: 1px solid #e5e7eb;
        background: #ffffff;
    }
    .metric-label {
        color: #6b7280;
        font-size: 0.9rem;
        margin-bottom: 0.2rem;
    }
    .metric-value {
        font-size: 1.8rem;
        font-weight: 800;
    }
    .stDataFrame thead tr th:nth-child(3),
    .stDataFrame thead tr th:nth-child(5),
    .stDataFrame tbody tr td:nth-child(3),
    .stDataFrame tbody tr td:nth-child(5) {
        text-align: center !important;
        justify-content: center !important;
    }
    .stDataFrame tbody td {
        padding: 0.35rem 0.5rem;
    }
    </style>
    """,
    unsafe_allow_html=True
)

st.title("📄 User Story Extractor")
st.markdown("Upload a `.docx` document with Epics, Stories, and Acceptance Criteria tables.")

uploaded_file = st.file_uploader("Upload Word Document", type=["docx"])
if uploaded_file:
    with st.spinner("Extracting stories and acceptance criteria..."):
        stories_df, ac_df = extract_user_stories_and_acs(uploaded_file)

    if stories_df.empty and ac_df.empty:
        st.warning("No user stories or acceptance criteria found in the document.")
    else:
        st.markdown('<div class="section-title">📊 Summary</div>', unsafe_allow_html=True)
        col1, col2, col3, col4 = st.columns(4)
        col1.metric("Total Epics", stories_df["Epic"].nunique() if not stories_df.empty else 0)
        col2.metric("Total Stories", len(stories_df))
        col3.metric("Total ACs", len(ac_df))
        avg_acs = (len(ac_df) / len(stories_df)) if (len(stories_df) > 0) else 0
        col4.metric("Avg ACs/Story", f"{avg_acs:.2f}")

        st.markdown('<div class="section-title">📑 Details</div>', unsafe_allow_html=True)
        tab1, tab2 = st.tabs(["📖 Story Details", "✅ Acceptance Criteria"])

        # ----- Tab 1: Story Details -----
        with tab1:
            f1, f2 = st.columns([1, 2])
            with f1:
                selected_epic = st.selectbox("Filter by Epic", ["All"] + (sorted(stories_df["Epic"].unique()) if not stories_df.empty else []))
            with f2:
                keyword = st.text_input("Search Title")

            filtered_df = stories_df.copy()
            if not stories_df.empty:
                if selected_epic != "All":
                    filtered_df = filtered_df[filtered_df["Epic"] == selected_epic]
                if keyword:
                    filtered_df = filtered_df[filtered_df["Story Title"].str.contains(keyword, case=False, na=False)]

            st.dataframe(filtered_df, use_container_width=True, hide_index=True)

            st.markdown('<div class="section-title">📥 Export Stories</div>', unsafe_allow_html=True)
            col1, col2 = st.columns(2)
            csv = filtered_df.to_csv(index=False).encode("utf-8-sig")
            with col1:
                st.download_button("Download CSV", csv, "stories.csv", "text/csv", use_container_width=True)
            excel_io = io.BytesIO()
            with pd.ExcelWriter(excel_io, engine="xlsxwriter") as writer:
                filtered_df.to_excel(writer, index=False, sheet_name="Stories")
            with col2:
                st.download_button("Download Excel", excel_io.getvalue(), "stories.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)

        # ----- Tab 2: Acceptance Criteria (only the requested columns) -----
        with tab2:
            g1, g2, g3 = st.columns([1, 1, 2])
            with g1:
                ac_epic = st.selectbox("Epic", ["All"] + (sorted(ac_df["Epic"].unique()) if not ac_df.empty else []))
            with g2:
                story_options = ["All"]
                if not ac_df.empty:
                    ac_scope = ac_df if ac_epic == "All" else ac_df[ac_df["Epic"] == ac_epic]
                    story_options += list(sorted(ac_scope["Story ID"].unique()))
                ac_story = st.selectbox("Story ID", story_options)
            with g3:
                ac_keyword = st.text_input("Search AC (Scenario / AC #)")

            ac_filtered = ac_df.copy()
            if not ac_df.empty:
                if ac_epic != "All":
                    ac_filtered = ac_filtered[ac_filtered["Epic"] == ac_epic]
                if ac_story != "All":
                    ac_filtered = ac_filtered[ac_filtered["Story ID"] == ac_story]
                if ac_keyword:
                    mask = (
                        ac_filtered["Scenario"].str.contains(ac_keyword, case=False, na=False) |
                        ac_filtered["AC #"].astype(str).str.contains(ac_keyword, case=False, na=False)
                    )
                    ac_filtered = ac_filtered[mask]

            preview_cols = ["Module","Epic","Story ID","Story Title","AC #","Scenario"]
            st.dataframe(ac_filtered[preview_cols], use_container_width=True, hide_index=True)

            st.markdown('<div class="section-title">📥 Export Acceptance Criteria</div>', unsafe_allow_html=True)
            h1, h2 = st.columns(2)
            ac_csv = ac_filtered[preview_cols].to_csv(index=False).encode("utf-8-sig")
            with h1:
                st.download_button("Download AC CSV", ac_csv, "acceptance_criteria.csv", "text/csv", use_container_width=True)
            ac_excel = io.BytesIO()
            with pd.ExcelWriter(ac_excel, engine="xlsxwriter") as writer:
                ac_filtered[preview_cols].to_excel(writer, index=False, sheet_name="Acceptance Criteria")
            with h2:
                st.download_button("Download AC Excel", ac_excel.getvalue(), "acceptance_criteria.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)
else:
    st.info("Please upload a Word document to extract user stories and acceptance criteria.")
